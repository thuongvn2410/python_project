import psycopg2
import random
import json
from docx import Document

# 🔹 Hàm kết nối PostgreSQL
def get_connection():
    return psycopg2.connect(
        dbname="",      
        user="",      
        password="",    
        host="",
        port=""
    )

# 🔹 Hàm tạo exam và các version
def create_exam(subject_id, code, title, duration_minutes, num_questions, version_codes):
    conn = get_connection()
    cur = conn.cursor()

    # 1) Tạo exam
    cur.execute("""
        INSERT INTO exams(subject_id, code, title, duration_minutes, num_questions)
        VALUES (%s,%s,%s,%s,%s)
        RETURNING id
    """, (subject_id, code, title, duration_minutes, num_questions))
    exam_id = cur.fetchone()[0]

    # 2) Lấy câu hỏi ngẫu nhiên
    cur.execute("""
        SELECT id, question, mix_choices
        FROM questions
        WHERE subject_id=%s
        ORDER BY RANDOM()
        LIMIT %s
    """, (subject_id, num_questions))
    questions = cur.fetchall()

    # 3) Tạo các version
    for v_code in version_codes:
        shuffle_seed = random.randint(1000, 9999)
        cur.execute("""
            INSERT INTO exam_versions(exam_id, version_code, shuffle_seed)
            VALUES (%s,%s,%s)
            RETURNING id
        """, (exam_id, v_code, shuffle_seed))
        ev_id = cur.fetchone()[0]

        rnd = random.Random(shuffle_seed)

        # 4) Xử lý từng câu hỏi
        for pos, (q_id, q_text, mix) in enumerate(questions, start=1):
            cur.execute("""
                SELECT id, content, position
                FROM choices
                WHERE question_id=%s
                ORDER BY position
            """, (q_id,))
            choices = cur.fetchall()

            if mix == 1:
                rnd.shuffle(choices)  # xáo trộn thứ tự
            order = [c[0] for c in choices]  # list id choice sau khi trộn

            cur.execute("""
                INSERT INTO exam_version_questions(exam_version_id, question_id, question_position, choice_order_json)
                VALUES (%s,%s,%s,%s)
            """, (ev_id, q_id, pos, json.dumps(order)))

        # 5) Xuất ra file Word
        export_exam_to_docx(cur, ev_id, f"Exam_{code}_{v_code}.docx")

    conn.commit()
    cur.close()
    conn.close()


# 🔹 Xuất ra file Word
def export_exam_to_docx(cur, exam_version_id, filename):
    doc = Document()

    # Lấy thông tin exam & version
    cur.execute("""
        SELECT e.title, e.duration_minutes, ev.version_code
        FROM exams e
        JOIN exam_versions ev ON e.id=ev.exam_id
        WHERE ev.id=%s
    """, (exam_version_id,))
    title, duration, version_code = cur.fetchone()

    doc.add_heading(title, level=1)
    doc.add_paragraph(f"Mã đề: {version_code} | Thời gian: {duration} phút\n")

    # Lấy câu hỏi đã snapshot
    cur.execute("""
        SELECT evq.question_id, evq.question_position, evq.choice_order_json, q.question
        FROM exam_version_questions evq
        JOIN questions q ON evq.question_id=q.id
        WHERE evq.exam_version_id=%s
        ORDER BY evq.question_position
    """, (exam_version_id,))
    q_rows = cur.fetchall()

    for pos, (q_id, q_pos, order_json, q_text) in enumerate(q_rows, start=1):
        doc.add_paragraph(f"Câu {pos}: {q_text}")
        order = json.loads(order_json)

        for cid in order:
            cur.execute("SELECT content FROM choices WHERE id=%s", (cid,))
            choice_text = cur.fetchone()[0]
            doc.add_paragraph(f"- {choice_text}", style="List Bullet")

    doc.save(filename)
    print(f"✅ Đã xuất đề {filename}")


# =============================
# Ví dụ chạy
# =============================
if __name__ == "__main__":
    create_exam(
        subject_id=1,
        code="MATH-MID-2025",
        title="ĐỀ THI GIỮA KỲ MÔN TOÁN",
        duration_minutes=60,
        num_questions=10,
        version_codes=["101", "102"]
    )
